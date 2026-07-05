#!/usr/bin/env python3
"""
export_project_docx_fixed.py

Safer .docx exporter for whole project, handles NULL/control bytes and large files.
Exports into multiple files (1.docx, 2.docx, etc.) capped at ~200 pages per file.
Also includes content from .txt files found in the project.

Usage:
  pip install python-docx
  python export_project_docx_fixed.py . --out-dir ./exports

Options:
  --max-size       max bytes per file to include contents, 0 to include all (default 200000)
  --skip           comma separated dir names to skip (default: .git,node_modules,venv,__pycache__)
  --include-hidden include dotfiles
  --out-dir        Directory to save the generated 1.docx, 2.docx files (default: current dir)
"""

import os
import sys
import argparse
import datetime
from datetime import timezone
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SKIP_DEFAULT = {'.git', 'node_modules', '__pycache__', '.venv', 'venv', '.ipynb_checkpoints'}

# Allowed control characters, keep newline and tab
ALLOWED_CONTROLS = {9, 10, 13}

# Heuristic for page counting
LINES_PER_PAGE = 55
PAGES_PER_DOC = 200
MAX_LINES_PER_DOC = LINES_PER_PAGE * PAGES_PER_DOC

# File extensions to include as text (readable files)
TEXT_EXTENSIONS = {
    # Plain text
    '.txt', '.log', '.csv', '.tsv', '.ini', '.cfg', '.conf', '.env',
    
    # Documentation
    '.md', '.markdown', '.rst', '.tex', '.ltx', '.adoc', '.asciidoc',
    '.org', '.wiki', '.pod',
    
    # Programming languages
    '.py', '.pyw', '.pyi', '.pyx', '.pxd', '.pxi',  # Python
    '.js', '.jsx', '.ts', '.tsx', '.mjs', '.cjs',   # JavaScript/TypeScript
    '.java', '.class', '.jar',                      # Java
    '.c', '.cpp', '.cxx', '.cc', '.h', '.hpp', '.hxx', '.hh',  # C/C++
    '.cs', '.csx',                                  # C#
    '.go',                                          # Go
    '.rs',                                          # Rust
    '.rb', '.rbw', '.gemspec',                      # Ruby
    '.php', '.php3', '.php4', '.php5', '.php7', '.phps', '.phpt',  # PHP
    '.pl', '.pm', '.t', '.pod',                    # Perl
    '.lua',                                         # Lua
    '.r', '.rmd',                                   # R
    '.swift',                                       # Swift
    '.kt', '.kts',                                  # Kotlin
    '.scala', '.sc',                                # Scala
    '.dart',                                        # Dart
    '.el', '.lisp', '.cl',                         # Lisp
    '.erl', '.hrl',                                 # Erlang
    '.ex', '.exs',                                  # Elixir
    '.cr',                                          # Crystal
    '.nim',                                         # Nim
    '.zig',                                         # Zig
    
    # Web technologies
    '.html', '.htm', '.xhtml', '.shtml',            # HTML
    '.css', '.scss', '.sass', '.less', '.styl',     # CSS
    '.xml', '.xsd', '.xsl', '.xslt', '.dtd',       # XML
    '.json', '.jsonc', '.jsonl',                    # JSON
    '.yaml', '.yml',                                 # YAML
    '.toml',                                        # TOML
    '.graphql', '.gql',                             # GraphQL
    '.vue', '.svelte',                              # Frontend frameworks
    
    # Shell scripts
    '.sh', '.bash', '.zsh', '.fish', '.ps1', '.psm1', '.psd1',
    '.bat', '.cmd', '.com',
    
    # Config files
    '.gitignore', '.dockerignore', '.editorconfig', 
    '.prettierrc', '.eslintrc', '.babelrc', '.npmrc',
    '.yarnrc', '.flake8', '.pylintrc', '.mypy.ini',
    '.coveragerc', '.stylelintrc', '.pre-commit-config.yaml',
    
    # Database
    '.sql', '.sqlite', '.db.sql',
    
    # Makefiles
    'Makefile', 'makefile', 'CMakeLists.txt',
    '.cmake', '.mk',
    
    # Other
    '.proto', '.thrift',                           # Protocol buffers
    '.wsdl', '.wadl',                              # Web services
    '.rss', '.atom',                               # Feeds
    '.svg',                                        # SVG (XML-based)
    '.csv', '.tsv',                                # Data
    '.diff', '.patch',                             # Patches
    '.license', '.copying',                        # License files
}

# Files to always include by name (even without extension)
TEXT_FILENAMES = {
    'Makefile', 'makefile', 'CMakeLists.txt',
    'Dockerfile', 'dockerfile',
    'README', 'readme', 'README.md', 'readme.md',
    'LICENSE', 'license', 'COPYING', 'copying',
    'CHANGELOG', 'changelog', 'CHANGES', 'changes',
    'CONTRIBUTING', 'contributing',
    'AUTHORS', 'authors',
    'TODO', 'todo', 'FIXME', 'fixme',
    '.gitignore', '.dockerignore', '.editorconfig',
    'package.json', 'package-lock.json', 'yarn.lock',
    'requirements.txt', 'Pipfile', 'Pipfile.lock',
    'setup.py', 'setup.cfg', 'pyproject.toml',
}

def is_probably_binary(path, blocksize=1024):
    try:
        with open(path, 'rb') as f:
            block = f.read(blocksize)
            if not block:
                return False
            # if NUL byte present, treat as binary
            if b'\x00' in block:
                return True
            # heuristics: high ratio of non-printable bytes
            textchars = bytearray({7,8,9,10,12,13,27} | set(range(0x20, 0x100)))
            nontext = block.translate(None, textchars)
            return float(len(nontext)) / max(1, len(block)) > 0.30
    except Exception:
        return True

def is_text_file(filename):
    """Check if file should be treated as text based on extension or name."""
    # Check by exact filename match
    basename = os.path.basename(filename)
    if basename in TEXT_FILENAMES:
        return True
    
    # Check by extension
    ext = os.path.splitext(filename)[1].lower()
    return ext in TEXT_EXTENSIONS

def sanitize_for_xml(s):
    # remove NULL and control characters except newline and tab
    out_chars = []
    for ch in s:
        code = ord(ch)
        if code == 0:
            continue
        if (0 <= code < 32) and (code not in ALLOWED_CONTROLS):
            # replace with unicode replacement char or skip
            out_chars.append('\uFFFD')
        else:
            out_chars.append(ch)
    return ''.join(out_chars)

def make_tree_lines(root, skip, include_hidden):
    lines = []
    for dirpath, dirs, files in os.walk(root):
        # filter
        if not include_hidden:
            dirs[:] = [d for d in dirs if not d.startswith('.')]
        rel = os.path.relpath(dirpath, root)
        depth = 0 if rel == '.' else rel.count(os.sep) + 1
        indent = '  ' * depth
        name = os.path.basename(dirpath) if rel != '.' else os.path.basename(os.path.abspath(root))
        if name in skip:
            dirs[:] = []
            continue
        lines.append(f"{indent}- {name}/")
        for f in sorted(files):
            if not include_hidden and f.startswith('.'):
                continue
            # Mark text files differently in the tree
            if is_text_file(f):
                lines.append(f"{indent}  - 📄 {f}")
            else:
                lines.append(f"{indent}  - 📁 {f}")
        dirs[:] = [d for d in dirs if d not in skip]
    return lines

class MultiDocManager:
    """Manages splitting output across multiple docx files based on line limits."""
    def __init__(self, output_dir, root_path):
        self.output_dir = output_dir
        self.root_path = root_path
        self.doc_index = 1
        self.current_lines = 0
        self.doc = Document()
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        self._init_doc()

    def _init_doc(self):
        self.doc.add_heading(f'Project export - Part {self.doc_index}', level=1)
        self.doc.add_paragraph(f'export time: {datetime.datetime.now(timezone.utc).isoformat()} UTC')
        self.doc.add_paragraph(f'root path: {self.root_path}')
        self.current_lines += 5 # Account for headings/meta

    def save_current_and_start_new(self):
        out_path = os.path.join(self.output_dir, f"{self.doc_index}.docx")
        self.doc.save(out_path)
        print(f"Wrote {out_path} (Reached ~{PAGES_PER_DOC} pages limit)")
        
        self.doc_index += 1
        self.doc = Document()
        self.current_lines = 0
        self._init_doc()

    def check_capacity(self, lines_to_add):
        if self.current_lines + lines_to_add >= MAX_LINES_PER_DOC:
            self.save_current_and_start_new()

    def add_heading(self, text, level=1):
        self.check_capacity(3) # Headings take up extra vertical space
        self.doc.add_heading(text, level=level)
        self.current_lines += 3

    def add_paragraph(self, text):
        # Estimate wrapped lines (assuming ~80 chars per line for standard width)
        lines_estimate = max(1, len(text) // 80 + 1)
        self.check_capacity(lines_estimate)
        self.doc.add_paragraph(text)
        self.current_lines += lines_estimate

    def add_page_break(self):
        self.check_capacity(LINES_PER_PAGE) 
        self.doc.add_page_break()
        # Fast forward the line count to the next multiple of LINES_PER_PAGE
        remainder = self.current_lines % LINES_PER_PAGE
        if remainder != 0:
            self.current_lines += (LINES_PER_PAGE - remainder)

    def add_code_line(self, text):
        lines_estimate = max(1, len(text) // 80 + 1)
        self.check_capacity(lines_estimate)
        
        p = self.doc.add_paragraph()
        # Remove paragraph spacing to make it look like tight code
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        r = p.add_run()
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        r.font.size = Pt(9)
        r.add_text(text)
        self.current_lines += lines_estimate

    def add_code_block(self, text):
        lines = text.splitlines()
        if not lines:
            self.add_code_line('')
            return
        for ln in lines:
            self.add_code_line(ln)

    def finish(self):
        if self.current_lines > 5: # Only save if we actually wrote content beyond headers
            out_path = os.path.join(self.output_dir, f"{self.doc_index}.docx")
            self.doc.save(out_path)
            print(f"Wrote {out_path}")

def read_file_content(path, max_size):
    """Read file content safely with proper encoding detection."""
    try:
        size = os.path.getsize(path)
        if max_size and size > max_size:
            return None, f"file skipped, larger than {max_size} bytes"
    except OSError:
        return None, "could not read file size"
    
    # First try UTF-8
    try:
        with open(path, 'r', errors='replace', encoding='utf-8') as fh:
            raw = fh.read()
            return raw, None
    except Exception:
        pass
    
    # Try latin-1 fallback
    try:
        with open(path, 'r', errors='replace', encoding='latin-1') as fh:
            raw = fh.read()
            return raw, None
    except Exception:
        return None, "could not read file content"

def main():
    parser = argparse.ArgumentParser(description='Safer project -> docx exporter (multi-file)')
    parser.add_argument('path', nargs='?', default='.')
    parser.add_argument('--out-dir', default='.', help='Directory to save the numbered docx files')
    parser.add_argument('--skip', default=','.join(sorted(SKIP_DEFAULT)))
    parser.add_argument('--max-size', type=int, default=200000)
    parser.add_argument('--include-hidden', action='store_true')
    parser.add_argument('--include-binary', action='store_true', 
                       help='Include binary files (skipped by default)')
    args = parser.parse_args()

    root = os.path.abspath(args.path)
    skip_set = set([x.strip() for x in args.skip.split(',') if x.strip()])

    doc_manager = MultiDocManager(output_dir=args.out_dir, root_path=root)

    # Add file tree
    doc_manager.add_heading('File tree', level=2)
    doc_manager.add_paragraph('📄 = Text file (content included)')
    doc_manager.add_paragraph('📁 = Binary file (content not included)')
    doc_manager.add_paragraph('')
    for ln in make_tree_lines(root, skip_set, args.include_hidden):
        doc_manager.add_paragraph(ln)
        
    doc_manager.add_page_break()
    doc_manager.add_heading('Files with contents', level=2)

    # Walk through all files
    for dirpath, dirs, files in os.walk(root):
        base = os.path.basename(dirpath)
        if base in skip_set:
            dirs[:] = []
            continue
        if not args.include_hidden:
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            files = [f for f in files if not f.startswith('.')]
            
        for fname in sorted(files):
            rel = os.path.relpath(os.path.join(dirpath, fname), root)
            path = os.path.join(dirpath, fname)
            
            try:
                size = os.path.getsize(path)
            except OSError:
                continue
                
            # Determine if this is a text file
            is_text = is_text_file(fname)
            
            # Add file header
            file_type = "📄 TEXT" if is_text else "📁 BINARY"
            doc_manager.add_heading(f"{file_type}: {rel}  (size: {size} bytes)", level=3)
            
            # Skip binary files if not explicitly included
            if not is_text and not args.include_binary:
                doc_manager.add_paragraph('Binary file (skipped). Use --include-binary to include.')
                continue
            
            # Read and include content for text files
            if is_text:
                content, error = read_file_content(path, args.max_size)
                if error:
                    doc_manager.add_paragraph(error)
                    continue
                if content is not None:
                    text = sanitize_for_xml(content)
                    doc_manager.add_code_block(text)
                else:
                    doc_manager.add_paragraph('File is empty or could not be read')
            else:
                # Binary file - show some info but don't include full content
                doc_manager.add_paragraph(f'Binary file: {fname}')
                doc_manager.add_paragraph(f'File type: {os.path.splitext(fname)[1] or "unknown"}')
                doc_manager.add_paragraph(f'Size: {size} bytes')
                doc_manager.add_paragraph('Full binary content not included in export.')
            
    try:
        doc_manager.finish()
        print('Export complete!')
        print(f'Files saved in: {os.path.abspath(args.out_dir)}')
        print('\nIncluded file types:')
        print('  - All text files: .txt, .md, .py, .js, .jsx, .ts, .tsx, .java, .c, .cpp, etc.')
        print('  - Web files: .html, .css, .scss, .json, .xml, .yaml, .toml')
        print('  - Config files: package.json, requirements.txt, Dockerfile, etc.')
        print('  - Binary files are listed but content is not included (use --include-binary to include)')
    except Exception as e:
        print('Error saving docx:', e, file=sys.stderr)
        sys.exit(2)

if __name__ == '__main__':
    main()